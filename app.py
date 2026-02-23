import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

st.set_page_config(page_title="劇本殺 A5 自動排版工具", page_icon="🎭")

st.title("🎭 劇本殺 A5 自動排版工具")
st.write("上傳純文字劇本，AI 將自動辨識懸疑反轉、強制分頁，並直接輸出 A5 格式的 Word 檔。")

# 讓使用者輸入自己的 Gemini API Key (也可以設定在 Streamlit 後台隱藏起來)
api_key = st.text_input("請輸入你的 Gemini API Key", type="password")

uploaded_file = st.file_uploader("上傳繁化好的劇本純文字檔 (.txt)", type=["txt"])

if st.button("開始全自動排版") and uploaded_file and api_key:
    with st.spinner("AI 正在閱讀劇本並構思排版節奏，請稍候..."):
        try:
            # 1. 讀取文本
            script_text = uploaded_file.read().decode("utf-8")
            
            # 2. 呼叫 Gemini AI 進行標籤化
            genai.configure(api_key=api_key)
            # 使用最新的模型
            model = genai.GenerativeModel('gemini-2.5-pro') 
            
            prompt = """
            你是一個專業的劇本殺排版編輯。請閱讀以下劇本，並在適當的地方插入特定標籤：
            1. 遇到一般章節結束、或需要停頓的地方，在該行獨立加上 [PAGE_BREAK]
            2. 遇到「未經主持人允許請勿翻開下一頁」等警告語，將該句獨立一行，並加上 [WARNING] 標籤
            3. 遇到情緒張力極高、重大反轉的單一關鍵句（例如：「原來，兇手就是你自己。」），將該句獨立一行，並加上 [SINGLE_SENTENCE] 標籤
            
            請直接輸出加上標籤後的完整劇本，不要加入任何其他的問候語或解釋。
            劇本內容如下：\n\n
            """
            
            response = model.generate_content(prompt + script_text)
            marked_text = response.text
            
            st.success("AI 處理完成！正在生成 A5 Word 檔案...")
            
            # 3. 使用 python-docx 生成 A5 Word 檔
            doc = Document()
            
            # 設定為 A5 尺寸 (148mm x 210mm)
            section = doc.sections[0]
            section.page_width = Mm(148)
            section.page_height = Mm(210)
            
            # 逐行解析 AI 處理後的文字
            for line in marked_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                    
                if '[PAGE_BREAK]' in line:
                    doc.add_page_break()
                elif '[WARNING]' in line:
                    doc.add_page_break()
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line.replace('[WARNING]', '').strip())
                    run.bold = True
                    run.font.size = Pt(16)
                    doc.add_page_break()
                elif '[SINGLE_SENTENCE]' in line:
                    doc.add_page_break()
                    # 插入幾個空白段落讓文字大致在頁面中間 (製造留白感)
                    for _ in range(5):
                        doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line.replace('[SINGLE_SENTENCE]', '').strip())
                    run.bold = True
                    run.font.size = Pt(14)
                    doc.add_page_break()
                else:
                    # 一般內文
                    doc.add_paragraph(line)
            
            # 4. 將寫好的 Word 存入記憶體，讓使用者下載
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="📥 點擊下載排版好的 A5 劇本",
                data=bio.getvalue(),
                file_name="劇本殺_A5排版完成.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.balloons() # 慶祝特效
            
        except Exception as e:
            st.error(f"發生錯誤：{e}")
